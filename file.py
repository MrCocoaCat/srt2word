# -*- coding: utf-8 -*-
import os, time
from docx import Document
import chardet


def srt2list(filepath):
    with open(filepath, mode='rb') as f:
        result = chardet.detect(f.read())
    # print(result['encoding'])
    # f1 = open(filepath, encoding=result['encoding'])
    # temp = f1.read()
    with open(filepath, encoding=result['encoding']) as f1:
        c = list(f1)
        a = [i.rstrip() for i in c]
    # print(a)
    b = [a[i:i + 4] for i in range(0, len(a), 4)]
    # print(b)
    return b


def cover(filepath, out_dirpath):
    print('Run task %s (%s)...' % (filepath, os.getpid()))
    start = time.time()

    if os.path.isfile(filepath):
        print(filepath)
        filename = os.path.basename(filepath)
        print(filename)
        name, suffix = os.path.splitext(filename)
        if suffix == ".srt":
            filename_doc = name + ".doc"
            filename_doc_path = os.path.join(out_dirpath, filename_doc)
            if os.path.isfile(filename_doc_path):
                os.remove(filename_doc_path)

            b = srt2list(filepath)

            document = Document()
            table = document.add_table(rows=0, cols=4, style="Table Grid")
            ### 判断
            # for line in f1.readlines():  # 遍历srt字幕文件
            for j in range(0, len(b)):
                cells = table.add_row().cells
                cells[0].text = str(b[j][0])
                cells[1].text = str(b[j][1])
                cells[2].text = str(b[j][2])
                cells[3].text = str(b[j][3])
            # paragraph = document.add_paragraph(temp)
            document.save(filename_doc_path)
    end = time.time()
    print('Task %s runs %0.2f seconds.' % (name, (end - start)))
    return filepath


def word_to_srt(filepath, out_dirpath):
    print('Run task %s (%s)...' % (filepath, os.getpid()))
    start = time.time()

    if os.path.isfile(filepath):
        print(filepath)
        filename = os.path.basename(filepath)
        print(filename)
        name, suffix = os.path.splitext(filename)
        if suffix == ".docx":
            filename_srt = name + ".srt"
            filename_srt_path = os.path.join(out_dirpath, filename_srt)
            if os.path.isfile(filename_srt_path):
                os.remove(filename_srt_path)
            document = Document(filename)
           # table = document.table()
           # print(document)
            ### 判断

    end = time.time()
    print('Task %s runs %0.2f seconds.' % (filepath, (end - start)))
    return filepath


if __name__ == '__main__':
    word_to_srt(filepath="C:\\Users\\54347\\Desktop\\1.docx",
                out_dirpath="C:\\Users\\54347\\Desktop")
